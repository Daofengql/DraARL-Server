from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "软著材料"
ASSETS = OUT / "assets"
TODAY = "2026年8月7日"
SOFTWARE = "DraARL数字无线电实时通信与设备管理平台"
VERSION = "V2.0"


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(run, field):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_document_defaults(doc: Document, header_title: str):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [("Title", 22, (20, 51, 83)), ("Heading 1", 15, (20, 51, 83)), ("Heading 2", 12, (38, 72, 104)), ("Heading 3", 10.5, (62, 62, 62))]:
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run(header_title)
    set_run_font(r, name="微软雅黑", size=8.5, color=(100, 100, 100))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("DraARL | 软件著作权登记材料草稿 | 第 ")
    set_run_font(r, name="微软雅黑", size=8, color=(120, 120, 120))
    add_field(footer.add_run(), "PAGE")
    r = footer.add_run(" 页")
    set_run_font(r, name="微软雅黑", size=8, color=(120, 120, 120))


def add_title_page(doc: Document, subtitle: str, note: str):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DraARL")
    set_run_font(r, name="微软雅黑", size=30, bold=True, color=(14, 104, 160))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SOFTWARE)
    set_run_font(r, name="微软雅黑", size=20, bold=True, color=(20, 51, 83))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, name="微软雅黑", size=15, bold=True, color=(38, 72, 104))
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"申报建议版本：{VERSION}")
    set_run_font(r, name="微软雅黑", size=12, color=(70, 70, 70))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"编制日期：{TODAY}")
    set_run_font(r, name="微软雅黑", size=11, color=(90, 90, 90))
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(note)
    set_run_font(r, name="微软雅黑", size=9.5, color=(160, 80, 0))
    doc.add_page_break()


def add_notice(doc, text, fill="FFF2CC"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_text(cell, text, size=9.5, color=(100, 60, 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_para(doc, text, bold_prefix=None, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, size=size, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=size, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color)
    return p


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_run_font(r, size=10.2)


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths=None, header_fill="D9EAF7"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        set_cell_text(header.cells[i], value, bold=True, color=(20, 51, 83), size=9.2)
        shade(header.cells[i], header_fill)
        if widths:
            header.cells[i].width = Inches(widths[i])
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=9.1)
            if widths:
                row.cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_image_if_exists(doc, path: Path, width=6.2, caption=None):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(caption)
            set_run_font(r, size=8.5, color=(100, 100, 100))
        return True
    add_notice(doc, f"图示文件未生成：{path.name}。不影响文字说明，提交前可补入实际系统截图。", "FCE4D6")
    return False


def make_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)
    overview = ROOT / "docs" / "assets" / "images" / "draarl-manual-overview.png"
    if overview.exists():
        shutil.copy2(overview, ASSETS / "draarl-manual-overview.png")
    magick = shutil.which("magick")
    names = ["overall-architecture", "udp-server-architecture", "auth-architecture", "voice-routing", "group-link-architecture", "cache-architecture", "device-binding-flow", "ota-flow"]
    if not magick:
        return
    for name in names:
        source = ROOT / "docs" / "assets" / "diagrams" / f"{name}.svg"
        target = ASSETS / f"{name}.png"
        if not source.exists() or target.exists():
            continue
        try:
            subprocess.run([magick, "-background", "white", "-density", "180", str(source), str(target)], check=True, capture_output=True, timeout=30)
        except (OSError, subprocess.SubprocessError):
            pass


def build_application_doc():
    doc = Document()
    set_document_defaults(doc, "软件著作权登记申请材料（填报稿）")
    add_title_page(doc, "软件著作权登记申请信息填报稿", "内部准备稿｜不能替代中国版权保护中心在线申请表")
    add_notice(doc, "使用说明：正式申请应登录中国版权保护中心登记系统在线填报并以系统生成的申请表为准。本稿将代码仓库信息整理为可复制字段；所有带“待确认”的身份、日期、发表和权属字段，提交前必须由著作权人核实。")

    add_heading(doc, "一、建议申报方案", 1)
    add_para(doc, "建议以独立开发的软件应用申请计算机软件著作权登记，申报名称保持稳定、功能描述聚焦实时数字无线电通信与设备管理，不把第三方框架或数据库写成自主创新。建议证书软件名称为“DraARL数字无线电实时通信与设备管理平台”，版本号为“V2.0”。")
    add_table(doc, ["项目", "建议填写", "依据/核对"], [
        ["软件全称", SOFTWARE, "软件定位、站点配置和文档均使用 DraARL；名称以“平台”结尾。"],
        ["软件简称", "DraARL", "internal/common/site.go 与前端站点配置。"],
        ["版本号", VERSION, "仓库 VERSION 为 v2.0.0-alpha4；证书建议使用稳定申报口径，勿把 alpha 当作功能版本。"],
        ["软件分类", "应用软件／通信与网络应用软件", "最终以登记系统可选分类为准。"],
        ["开发方式", "原始取得／独立开发（待确认）", "提交记录显示单一开发者署名；仍需核对是否存在共同开发、委托或职务开发。"],
        ["权利范围", "全部", "如无转让、许可或限制，通常选择全部。"],
        ["软件著作权人", "【待填写：个人姓名或单位全称】", "必须与身份证/营业执照及权属证明完全一致。"],
        ["申报版本截止", "2026-08-07 仓库状态（待确认）", "当前工作区版本；正式提交时应固定源码快照和校验值。"],
    ], widths=[1.25, 2.4, 3.0])

    add_heading(doc, "二、申请表字段草案", 1)
    add_table(doc, ["字段", "填报草案", "需确认事项"], [
        ["软件名称", SOFTWARE, "确认是否需加入“软件”二字；名称应与说明书、源码文档一致。"],
        ["软件简称", "DraARL", "没有简称可留空；如填，全文保持一致。"],
        ["版本号", VERSION, "确认是否按 V2.0 申报；不要把 v2.0.0-alpha4 与 V2.0 混用。"],
        ["开发完成日期", "【待确认】", "代码首次提交为 2026-03-10，仅是开发记录，不等同法律完成日。"],
        ["首次发表日期", "【待确认：未发表／已发表及日期】", "GitHub、在线站点和文档是否构成公开发表，需按实际发布时间核实。"],
        ["首次发表地点/网址", "【待确认】", "如已发表，填写实际公开平台；不要仅凭仓库创建时间猜测。"],
        ["开发方式", "原始取得／独立开发（待确认）", "如有单位任务、委托合同或共同开发，应改为对应选项并准备证明。"],
        ["权利取得方式", "原始取得（待确认）", "如权利来自转让、继承或其他方式，必须改填并提供证明。"],
        ["权利范围", "全部（待确认）", "核对是否存在已许可、已转让或共同权利。"],
        ["著作权人", "【待填写】", "个人姓名/单位全称、证件号、地址、联系人等必须从证件复制。"],
        ["软件用途", "业余无线电及自研数字电台设备的实时通信、设备接入、群组协作和平台运维管理。", "正式系统若有字数限制，保留核心业务词。"],
        ["源程序量", "生产源代码约 91,516 行（343 个 Go/TypeScript/TSX/CSS 文件；不含测试、依赖、文档和构建产物）", "提交前重新固定快照统计，避免与源码文档页数矛盾。"],
        ["编程语言", "Go、TypeScript、TSX、CSS；配置和数据脚本含 YAML、SQL", "按登记系统字段填写主要语言，配置文件可在说明书中说明。"],
        ["开发环境", "Go 1.25、Node.js 20、React 19、TypeScript 5.9、Vite 7、Git；Windows/Linux 均可", "以实际开发机和最终源码快照为准。"],
        ["运行环境", "Linux/Windows/macOS；Go 服务、MySQL/MariaDB、Redis（可选）、MinIO（可选）；浏览器前端", "按实际交付部署方式删减。"],
    ], widths=[1.35, 3.65, 1.65])

    add_heading(doc, "三、软件功能概述（可粘贴至申请表）", 1)
    add_para(doc, "DraARL数字无线电实时通信与设备管理平台面向业余无线电和自研数字电台设备，提供账号注册登录、审核与权限控制、UDP实体设备接入、WebSocket幽灵设备接入、设备动态码绑定、设备参数配置同步、公开/私有群组及虚拟互联组管理、实时语音与文本通信、通信记录和通联日志、APRS位置服务、固件与客户端资源发布、管理员后台和中心/边缘节点互联等功能。系统以自研DraARLv1协议承载设备通信，支持Opus语音、半双工话权、跨群组及跨节点路由、会话恢复、消息历史查询和对象存储。")

    add_heading(doc, "四、技术特点与自主开发点（可粘贴至申请表/说明书）", 1)
    add_bullets(doc, [
        "采用自研 DraARLv1 二进制协议，统一承载心跳、认证、配置、语音和文本消息，并对设备字段、包类型和身份进行校验。",
        "采用按通信域计算的半双工 Speaker Lease/话权仲裁，同一转发域同一时刻只允许一个有效说话人，超时自动释放。",
        "采用中心/边缘 Type 0 节点互联，控制面使用 TLS，数据面使用 HMAC 信封、会话世代、重放窗口和路由投影版本。",
        "采用按通信域缓存的接收计划和按 Session 去重的 fan-out，支持同账号多端在线、单发多收和来源频道标记。",
        "采用设备动态码绑定、AES/bcrypt 设备凭据、JWT/refresh token、Origin 白名单、接口限流和管理员审计等安全措施。",
        "采用 MySQL 持久化、Redis 可选令牌存储、MinIO/本地对象存储和多级内存缓存，支持单机和分布式部署。",
    ])

    add_heading(doc, "五、功能模块清单", 1)
    add_table(doc, ["模块", "主要功能", "对应代码/文档依据"], [
        ["账号与认证", "密码登录、邮箱验证码、SSO、JWT/refresh token、审核和权限", "internal/auth、internal/middleware、docs/usage/02"],
        ["设备接入", "UDP设备认证、动态码绑定、设备型号/SSID、上线状态、配置同步", "internal/protocol、internal/udphub、docs/usage/04、Protocol.md"],
        ["实时通信", "Opus语音、文本消息、PTT、WebSocket、多频道接收和来源频道", "internal/udphub、pkg/websocket、www/src/pages/radio"],
        ["群组与互联", "公开/私有群组、群成员、虚拟互联组、跨群路由和循环防护", "internal/groupaccess、internal/handler、docs/节点互联协议.md"],
        ["通信记录", "发信记录、录音、趋势统计、通联日志、频道消息历史", "internal/udphub、internal/gormdb、docs/usage/05"],
        ["管理后台", "用户、设备、群组、中继台、节点、资源、固件、站点配置和操作日志", "www/src/pages/admin、internal/handler、docs/usage/06"],
        ["发布与运维", "配置模板、数据库迁移、缓存指标、健康检查、多平台构建和资源分发", "cmd/draarl、internal/config、.github/workflows、docs/usage/01/08/09"],
    ], widths=[1.2, 3.1, 2.25])

    add_heading(doc, "六、权属和版本核对", 1)
    add_notice(doc, "仓库 NOTICE 写明 Copyright (c) 2026 Daofengql，Git 提交记录显示单一开发者署名“道锋潜鳞”。这只能作为内部线索，不能替代证件和权属证明。申请人必须核对：代码作者、著作权人、许可证声明和实际申报主体是否为同一主体；如不一致，应准备职务开发、委托开发、权利转让或共同开发证明。", "FCE4D6")
    add_table(doc, ["核对项", "当前仓库线索", "提交前动作"], [
        ["代码作者", "Git 提交记录仅见“道锋潜鳞”署名", "以真实姓名和证件核对；如为网名，准备作者/权利人说明。"],
        ["许可证", "PolyForm Noncommercial 1.0.0", "保留 LICENSE、NOTICE；许可证不等于著作权登记证明。"],
        ["公开状态", "GitHub仓库、在线站点和文档链接存在", "确认首次公开日期；考虑专利新颖性风险。"],
        ["版本快照", "VERSION=v2.0.0-alpha4，当前为 alpha 预发布版本", "在登记前打 tag、保存源码压缩包和 SHA-256。"],
    ], widths=[1.3, 2.7, 2.55])

    add_heading(doc, "七、正式提交材料清单", 1)
    add_bullets(doc, [
        "中国版权保护中心登记系统在线填写并生成的《计算机软件著作权登记申请表》；本文件只是填报草案。",
        "著作权人身份证明：个人身份证正反面或单位营业执照/统一社会信用代码证件；主体名称必须一致。",
        "源程序鉴别材料：使用本目录的《源程序文档》，按系统当期要求核对页码、页眉、行数、连续页和隐去敏感信息。",
        "文档鉴别材料：使用《软件设计说明书》；如系统要求用户手册或操作说明，可在其基础上补充运行截图。",
        "如开发方式或权利取得方式不是独立原始开发，补充任务书、委托合同、共同开发协议、权利转让/继承证明等。",
        "提交前固定源码、说明书和版本号，保存压缩包、SHA-256、构建日志及可运行演示环境；不要把 JWT、AES、数据库密码或用户数据带入材料。",
    ])
    add_heading(doc, "八、官方入口与注意事项", 1)
    add_para(doc, "中国版权保护中心官网：https://www.ccopyright.com.cn/；登记系统入口以官网当前链接为准：https://register.ccopyright.com.cn/。官网在线表单、材料格式和收费/办理规则可能调整，提交时应以系统页面和当期办事指南为准。")
    add_para(doc, "本稿不构成法律意见，也不保证登记机关对软件名称、版本、权属或材料页数的最终审核结论。")
    path = OUT / "01-DraARL软件著作权登记申请信息填报稿.docx"
    doc.save(path)
    return path


def build_manual_doc():
    doc = Document()
    set_document_defaults(doc, "软件设计说明书（鉴别材料草稿）")
    add_title_page(doc, "软件设计说明书（鉴别材料草稿）", "可作为软件功能/技术文档基础｜提交前按实际运行版本补充截图")
    add_notice(doc, "文档定位：本说明书描述软件的功能组成、技术架构、关键流程、运行环境和安全设计，可作为软件鉴别材料的基础文档。它不替代官方申请表；图片来自仓库文档或结构图，正式提交前应检查是否需要补充实际界面截图。")

    add_heading(doc, "1. 软件概述", 1)
    add_para(doc, "DraARL数字无线电实时通信与设备管理平台（DraARL）是面向业余无线电和自研数字电台设备的实时通信与管理软件。软件采用 Go 后端和 React/TypeScript 前端，向设备提供 UDP DraARLv1 接入服务，向浏览器提供 HTTP API、WebSocket 实时通信和管理界面，并通过 MySQL、Redis、MinIO 等组件完成持久化、令牌、对象和音频数据管理。")
    add_image_if_exists(doc, ASSETS / "draarl-manual-overview.png", 6.2, "图1  DraARL 平台场景概览（仓库 docs/assets/images/draarl-manual-overview.png）")
    add_table(doc, ["属性", "内容"], [
        ["软件名称", SOFTWARE], ["申报版本", VERSION], ["软件类型", "服务器端通信与管理应用软件，含浏览器前端"], ["主要用户", "业余无线电用户、设备管理员、节点运维人员、客户端/设备开发者"], ["服务形态", "单机部署或中心/边缘分布式部署；前端可嵌入 Go 二进制"], ["核心协议", "DraARLv1（公开设备协议）和 Type 0 中心/边缘互联协议"],
    ], widths=[1.55, 4.9])

    add_heading(doc, "2. 建设目标与适用范围", 1)
    add_bullets(doc, [
        "为数字电台设备提供统一的身份认证、接入、状态上报、配置同步和固件发布能力。",
        "为多种实体设备、移动/桌面客户端和浏览器提供统一的群组语音、文本消息和通信记录能力。",
        "通过虚拟互联组和中心/边缘节点拓扑，支持跨群组、跨站点和跨地域实时转发。",
        "为平台管理员提供用户审核、设备/群组/节点、资源/固件、站点配置、指标和操作审计工具。",
    ])
    add_para(doc, "软件不直接实现无线射频硬件驱动和电台固件本体；软件负责服务器侧协议接入、路由、管理、记录和发布，硬件设备通过 DraARLv1 协议接入。")

    add_heading(doc, "3. 总体技术架构", 1)
    add_image_if_exists(doc, ASSETS / "overall-architecture.png", 6.2, "图2  DraARL 整体架构")
    add_table(doc, ["层次", "组成", "职责"], [
        ["客户端层", "ESP32设备、移动/桌面客户端、浏览器、桥接器", "采集/播放语音，发送文本，接收状态，发起设备配置或管理操作。"],
        ["网络接入层", "UDP Server、HTTP API、WebSocket、TLS控制面", "完成包接收、鉴权、限流、跨域/来源检查和节点控制连接。"],
        ["业务逻辑层", "认证、设备、群组、路由、语音、记录、APRS、资源、审核", "把协议事件转换为权限检查、路由投影、持久化和管理操作。"],
        ["数据存储层", "MySQL/MariaDB、Redis、MinIO/本地对象存储、内存缓存", "存储账号/设备/群组/记录，缓存令牌和路由，保存文件、音频、固件。"],
        ["外部服务", "APRS-IS、Keycloak、SMTP、OpenAI（可选）", "提供位置网络、单点登录、验证码邮件和可选智能服务。"],
    ], widths=[1.15, 2.5, 3.1])

    add_heading(doc, "4. 服务端与前端组成", 1)
    add_table(doc, ["组成", "主要目录/组件", "说明"], [
        ["服务入口", "cmd/draarl", "加载配置、初始化数据库/存储/缓存，启动 UDP、HTTP、APRS 和中心/边缘模式。"],
        ["HTTP/API", "internal/server、internal/handler、internal/middleware", "Gin 路由、REST API、认证中间件、权限、限流、静态前端服务。"],
        ["协议与转发", "internal/protocol、internal/udphub、internal/interconnect", "DraARLv1 编解码、设备运行态、语音 fan-out、Type 0 节点中继。"],
        ["数据访问", "internal/gormdb、internal/db", "GORM 模型/仓储与兼容原生 SQL，提供迁移、分页、索引和统计。"],
        ["公共能力", "pkg/jwt、pkg/websocket、pkg/storage、pkg/cache 等", "JWT、WebSocket协议、对象存储、缓存、加密和 GeoIP 等可复用能力。"],
        ["浏览器前端", "www/src", "React 路由、用户控制台、在线收发页、管理后台、资源/记录/设置页面。"],
    ], widths=[1.2, 2.45, 3.1])

    add_heading(doc, "5. 核心模块说明", 1)
    modules = [
        ("5.1 账号、认证与权限", "支持账号密码、邮箱验证码、Keycloak SSO、JWT/refresh token；通过注册审核、操作证审核、管理员角色、群组成员和设备所有权形成多层权限。refresh token 优先存 Redis，Redis 不可用时可降级到内存存储。"),
        ("5.2 设备接入与绑定", "普通 UDP 设备使用设备密码认证；新设备通过预检查、6 位动态码和 Web 端确认建立用户与设备关系；服务端保存设备型号、MAC、SSID、群组、状态和参数，支持配置读取、保存和同步。"),
        ("5.3 群组与虚拟互联", "提供公开群组、私有群组、成员管理和群主操作；管理员可创建虚拟互联组，把多个目标群组连接为通信域，并通过路由缓存、防循环和权限检查进行跨组转发。"),
        ("5.4 实时语音与文本", "UDP DraARLv1 和 WebSocket 共同进入实时路由；语音使用 Opus 16K 帧，文本消息与语音使用统一通信域，浏览器支持 PTT、文本输入、发送频道和多接收频道。"),
        ("5.5 通信记录与日志", "按发送和投递关系保存通信记录，支持音频缓冲/对象存储、文本历史、用户发信记录、统计趋势和 HAM 通联日志；频道历史按发送时投递快照和游标索引查询。"),
        ("5.6 管理后台", "管理用户、设备、群组、中继台、中心/边缘节点、资源、固件、客户端资源、SMTP、APRS、OpenAI、通信设置、缓存指标、UDP指标和操作日志。"),
        ("5.7 节点互联", "中心承载数据库、权威路由和权限；边缘在无数据库模式承载设备入口和本地 fan-out。TLS 控制面协商能力、同步路由快照/增量、续租设备短期会话；Type 0 UDP 信封中继跨边缘数据。"),
    ]
    for title, text in modules:
        add_heading(doc, title, 2)
        add_para(doc, text)

    add_heading(doc, "6. 关键业务流程", 1)
    add_heading(doc, "6.1 设备动态码绑定", 2)
    add_image_if_exists(doc, ASSETS / "device-binding-flow.png", 6.1, "图3  设备动态码绑定流程")
    add_bullets(doc, [
        "设备提交 MAC、用户名和设备密码进行预检查；服务端判断设备是否需要绑定。",
        "需要绑定时，设备申请一次性动态码并显示；动态码短期有效且单次使用。",
        "审核通过用户在设备管理页输入动态码，服务端校验所有权、设备状态、SSID 和群组规则。",
        "绑定成功后，服务端下发用户名、设备密码、SSID、群组和设备参数；设备确认后进入正常心跳和通信流程。",
    ])
    add_heading(doc, "6.2 实时语音转发", 2)
    add_image_if_exists(doc, ASSETS / "voice-routing.png", 6.1, "图4  语音路由与半双工仲裁")
    add_bullets(doc, [
        "设备发送 Opus 语音帧，入口先完成设备/Session 身份、包类型、权限和新鲜度检查。",
        "服务端根据群组及虚拟互联关系计算通信域，并尝试取得唯一话权；无话权或被禁发的帧不进入转发。",
        "服务端从缓存接收计划生成目标快照，按 Session 去重、排除来源，再分配到发送 worker。",
        "连续语音停止超过约 900ms 后释放话权；通信记录器按会话缓冲并异步写入对象存储/数据库。",
    ])
    add_heading(doc, "6.3 中心/边缘跨节点通信", 2)
    add_image_if_exists(doc, ASSETS / "udp-server-architecture.png", 6.1, "图5  UDP 服务与节点数据面")
    add_bullets(doc, [
        "边缘通过 TLS 控制面注册、鉴权并协商协议版本和功能位；中心为节点建立 NodeSessionID、KeyEpoch 和会话密钥。",
        "中心向边缘发送路由快照或增量；边缘以投影版本 ACK，落后时请求重同步，避免使用过期路由。",
        "设备上行在来源边缘本地 fan-out，同时以一份 RelayUpstream 发送中心；中心校验节点会话和权威路由，再向目标边缘发送 RelayDownstream。",
        "中心重启后通过设备短期会话和幽灵恢复票据重新确认会话，失败或超时的会话进入 fail-closed 清理。",
    ])

    add_heading(doc, "7. 协议与数据设计", 1)
    add_table(doc, ["协议/数据", "关键设计", "安全/一致性"], [
        ["DraARLv1", "固定头部 + 类型 + 设备字段 + DATA；支持心跳、认证、配置、Opus、文本", "校验版本、长度、类型、SSID、设备型号；转发前清理敏感字段。"],
        ["Type 0 节点信封", "节点 ID、协议版本、Session、MessageID、时间、Hop、Payload、HMAC-SHA256", "TLS 控制面建立密钥；会话世代、重放窗口、HMAC 和目标版本防篡改/重放。"],
        ["用户/设备/群组", "用户、设备、群组、成员、虚拟互联和设备配置模型", "数据库约束 + 权限仓储；设备所有权和群组权限在路由前检查。"],
        ["通信记录", "记录主体 + 发送时投递快照 + 消息类型/文本/音频对象", "按游标和类型索引查询；历史可见性不依赖当前拓扑回算。"],
        ["缓存计划", "按通信域构建不可变接收计划，路由变更时原子替换", "Session 去重、来源精确排除、缓存指标和过载完整帧淘汰。"],
    ], widths=[1.35, 3.2, 2.3])

    add_heading(doc, "8. 安全设计", 1)
    add_bullets(doc, [
        "账号安全：密码哈希、邮箱验证码限流、JWT 访问令牌和 HttpOnly refresh/ws cookie；禁止把令牌放入 WebSocket URL。",
        "设备安全：设备凭据 AES 加密存储，兼容历史 bcrypt 校验；动态码单次短期使用，设备 Session 有续租和超时。",
        "接口安全：CORS/Origin 白名单、安全响应头、登录/绑定/消息 API 限流、管理员/审核/群组权限中间件。",
        "节点安全：TLS 控制面、节点凭据轮换、HMAC 数据面、随机挑战绑定地址、MessageID 重放窗口、资源上限和 fail-closed。",
        "数据安全：敏感配置通过部署侧密钥注入；申请材料不得包含 JWT Secret、DeviceAuth AESKey、数据库密码、对象存储密钥或真实用户数据。",
    ])
    add_image_if_exists(doc, ASSETS / "auth-architecture.png", 6.1, "图6  认证架构")

    add_heading(doc, "9. 前端功能入口", 1)
    add_table(doc, ["页面/入口", "功能"], [
        ["/", "/about、/docs、/forum、/relays、/tools 等公共入口和站点信息。"],
        ["/login、/register、/forgot-password、/sso/callback", "密码/验证码登录、注册、找回密码和 SSO 回调。"],
        ["/dashboard、/profile", "用户状态、统计、个人资料、头像、邮箱/密码、操作证和设备密码。"],
        ["/devices、/groups", "设备绑定/配置/群组切换，以及公开/私有群组和成员管理。"],
        ["/radio", "浏览器幽灵设备的 PTT、文本、多接收频道、来源频道和在线状态。"],
        ["/comm-records/platform、/comm-records/logbook", "平台通信记录、趋势、音频和 HAM 通联日志。"],
        ["/admin/*", "管理员用户、设备、群组、节点、资源、固件、记录、设置和指标。"],
    ], widths=[2.6, 4.25])

    add_heading(doc, "10. 部署与运行", 1)
    add_para(doc, "软件可使用 Docker Compose 运行 MySQL、Redis、MinIO 和 DraARL，也可手工部署 Go 服务和 React 构建产物。中心模式启动完整 HTTP/API、WebSocket、UDP 和数据依赖；边缘模式使用 --edge 启动无数据库的设备入口、本地转发和 TLS/UDP 节点互联。前端可以以 Vite 开发服务器运行，也可以构建后嵌入 Go 二进制或由独立 Web 服务器托管。")
    add_table(doc, ["环境", "最低/建议配置"], [
        ["开发环境", "Go 1.25+、Node.js 20+、npm、Git；MySQL/MariaDB；Redis/MinIO 按功能启用。"],
        ["服务端操作系统", "Linux、Windows 或 macOS；生产建议 Linux。"],
        ["小规模部署", "2 vCPU、4 GiB 内存、SSD；MySQL 必需，Redis/MinIO 按功能选择。"],
        ["高并发转发验证", "仓库记录过 4 vCPU/8 GiB 和 8 vCPU/8 GiB 的 Linux Hyper-V 测试；实测数据不构成生产 SLA。"],
        ["默认端口", "UDP 60050、HTTP 9000、节点 TLS 控制面 60100（均可在 config.yaml 调整）。"],
    ], widths=[1.7, 5.15])

    add_heading(doc, "11. 性能验证摘要", 1)
    add_para(doc, "仓库测试文档记录：在 Ubuntu 22.04、4 vCPU/8 GiB、同机 MySQL/Redis/DraARL/压测端环境中，9000 台设备、20ms 连续 30 秒达到 449,950 pps、零丢包，平均延迟 9.45ms；中心加三个边缘、8 vCPU/8 GiB、20,000 个设备的 30ms 档约 667k pps 且零丢包。上述结果只说明特定回环和虚拟机条件下的实现能力，不应写成公网 SLA 或产品容量承诺。")
    add_table(doc, ["验证项目", "结果摘要", "说明"], [
        ["UDP fan-out", "按完整帧排队、writer 分片、过载淘汰最旧完整帧", "避免每目标 job 造成的队列碎片和尾部随机截断。"],
        ["幽灵多收", "按通信域缓存接收计划，Session 去重；缓存命中约 27-28ns（测试环境）", "测试基准用于比较实现方案，不是生产 SLA。"],
        ["频道历史", "投递快照 + group/type/time/id 复合索引 + 游标分页", "避免按当前拓扑猜测历史可见性。"],
        ["节点互联", "中心/边缘跨节点 UDP relay、TLS 控制和会话恢复测试", "部署前仍需在目标网络复测公网延迟和带宽。"],
    ], widths=[1.35, 3.15, 2.35])

    add_heading(doc, "12. 版权与材料边界", 1)
    add_notice(doc, "本说明书只描述本项目自主代码和软件功能。Go/React/Gin/GORM/数据库、Opus 编解码和其他第三方组件的版权仍归各自权利人；登记材料应避免把第三方代码、依赖目录、node_modules、前端 dist、真实密钥和用户数据作为本软件原创内容。", "FCE4D6")
    add_para(doc, "源码快照建议在提交前通过 Git tag 固定，并生成 SHA-256；说明书中的版本、模块和图片应与该快照一致。若登记主体不是 Git 记录中的开发者本人，必须同步准备权属证明。")
    path = OUT / "02-DraARL软件设计说明书.docx"
    doc.save(path)
    return path


def collect_source_lines() -> list[str]:
    files = []
    for base in [ROOT / "cmd", ROOT / "internal", ROOT / "pkg", ROOT / "www" / "src"]:
        if not base.exists():
            continue
        files.extend(p for p in base.rglob("*") if p.is_file() and p.suffix.lower() in {".go", ".ts", ".tsx", ".css"} and not p.name.endswith("_test.go"))
    files = sorted(files, key=lambda p: p.as_posix())
    lines: list[str] = []
    for path in files:
        rel = path.relative_to(ROOT).as_posix()
        lines.append(f"// FILE: {rel}")
        try:
            content = path.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        lines.extend(content.splitlines())
    return lines


def add_code_page(doc, lines: Sequence[str], page_no: int):
    if page_no > 1:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{SOFTWARE} {VERSION} 源程序鉴别材料 | 第 {page_no}/60 页")
    set_run_font(r, name="微软雅黑", size=8.5, bold=True, color=(20, 51, 83))
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(8.2)
        p.paragraph_format.keep_together = True
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=6.4, color=(25, 25, 25))


def build_source_doc():
    doc = Document()
    set_document_defaults(doc, "源程序鉴别材料（前后各30页）")
    doc.core_properties.title = f"{SOFTWARE}{VERSION}源程序鉴别材料"
    doc.core_properties.subject = "前后各30页，每页50行，共60页"
    section = doc.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)
    all_lines = collect_source_lines()
    if len(all_lines) < 3000:
        all_lines.extend(["// [SOURCE LINE PLACEHOLDER]"] * (3000 - len(all_lines)))
    chunks = [all_lines[:1500], all_lines[-1500:]]
    page = 1
    for group in chunks:
        for i in range(30):
            lines = group[i * 50 : (i + 1) * 50]
            add_code_page(doc, lines, page)
            page += 1
    path = OUT / "03-DraARL源程序鉴别材料-前后各30页.docx"
    doc.save(path)
    return path, len(all_lines)


def write_readme(source_line_count: int):
    release_version = (ROOT / "VERSION").read_text(encoding="utf-8").strip()
    text = f"""# DraARL 软著材料草稿\n\n生成日期：{TODAY}\n\n本目录是根据仓库当前代码和文档整理的软著准备材料，不是官方登记表。正式提交请以中国版权保护中心登记系统当前页面和办事指南为准。\n\n## 文件\n\n- `01-DraARL软件著作权登记申请信息填报稿.docx`：把软件名称、版本、环境、功能概述、技术特点和权属核对整理成可复制字段。\n- `02-DraARL软件设计说明书.docx`：功能、架构、流程、安全、部署和性能摘要，可作为文档鉴别材料基础。\n- `03-DraARL源程序鉴别材料-前后各30页.docx`：从生产代码中按前后各 30 页、每页 50 行的目标格式生成的源码节选。\n- `assets/`：由仓库现有图片和架构 SVG 转换的文档插图。\n\n## 当前代码依据\n\n- 版本文件：`VERSION` 当前为 `{release_version}`；申请稿建议以稳定口径 `V2.0` 填报。\n- 生产源代码统计：约 91,516 行、343 个 Go/TypeScript/TSX/CSS 文件；不含测试、依赖、文档和构建产物。\n- 源码鉴别文档使用合并源码 `{source_line_count}` 行作为节选来源，正式提交前应对固定 tag 重新生成。\n- Git 提交记录当前仅见单一署名“道锋潜鳞”，不能替代真实权属证明。\n\n## 提交前必须确认\n\n1. 著作权人是个人还是单位，证件/营业执照上的全称、证件号、地址和联系人。\n2. 开发方式是否确实为独立开发，是否存在职务开发、委托开发、共同开发或权利转让。\n3. 开发完成日期、首次发表状态、首次发表日期和网址。代码首次提交日 `2026-03-10` 只作开发记录，不自动等于法律上的完成日。\n4. 申报版本源码、设计说明书、前端页面和线上运行版本是否一致；建议固定 Git tag，保存源码包和 SHA-256。\n5. 源码/说明书是否含 JWT Secret、DeviceAuth AESKey、数据库/对象存储密码、邮箱密钥、用户资料或其他个人信息。\n6. 打印或转 PDF 后，复核源码是否为连续前后各 30 页、每页行数和页码是否满足登记系统当期要求；不足 60 页时按当期要求提交全部或调整排版。\n\n## 申请什么最合适\n\n- **优先：软件著作权登记。** 保护本项目源代码和文档的表达，适合作为项目权属、招投标、交付、维权和版本留痕材料。\n- **并行评估：`DraARL` 商标。** 软著不保护品牌名称，名称、Logo、官网和客户端品牌应另做商标检索，类别通常需要让代理人结合实际服务选择。\n- **谨慎评估：发明专利。** DraARLv1、跨节点低时延转发、话权仲裁等可能包含技术方案，但仓库/GitHub/在线站点已经公开，是否仍具新颖性必须先做专利检索；不要先继续公开细节再申请。\n- **持续管理：商业秘密和许可证。** 未公开的密钥、部署配置、运营数据和未发布实现可通过权限、保密和访问控制管理；现有 PolyForm Noncommercial 许可证和 NOTICE 需与实际著作权人保持一致。\n\n## 官方入口\n\n中国版权保护中心：<https://www.ccopyright.com.cn/>\n\n登记系统入口（以官网当前链接为准）：<https://register.ccopyright.com.cn/>\n\n"""
    (OUT / "00-README-申报建议与待确认事项.md").write_text(text, encoding="utf-8")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    make_assets()
    build_application_doc()
    build_manual_doc()
    _, count = build_source_doc()
    write_readme(count)
    print(f"Generated soft copyright materials in {OUT}")


if __name__ == "__main__":
    main()
