from __future__ import annotations

import hashlib
import re
import textwrap
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from generate_softcopyright_docs import (
    ROOT,
    SOFTWARE,
    VERSION,
    add_bullets,
    add_field,
    add_heading,
    add_notice,
    add_para,
    add_table,
    collect_source_lines,
    prevent_row_split,
    set_cell_text,
    set_document_defaults,
    set_repeat_table_header,
    set_run_font,
    shade,
)


DOCX_OUT = ROOT / "output" / "软著材料" / "官方要求版"
PDF_OUT = ROOT / "output" / "pdf" / "软著材料"
ACCESS_DATE = "2026-08-07"
OFFICIAL_GUIDE = "https://www.ccopyright.com.cn/index.php?optionid=1033"
OFFICIAL_FILES = "https://www.ccopyright.com.cn/index.php?optionid=1080"
OFFICIAL_FORM_HELP = "https://www.ccopyright.com.cn/index.php?optionid=1081"
OFFICIAL_SYSTEM = "https://register.ccopyright.com.cn/registration.html#/registerSoft"


def configure_page(doc: Document, title: str, body_font_size=10.5):
    set_document_defaults(doc, title)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(body_font_size)
    doc.core_properties.title = title
    doc.core_properties.subject = f"{SOFTWARE} {VERSION} 软件著作权登记材料"


def configure_identification_page(doc: Document, material_name: str, top=1.45, bottom=1.25, side=1.15):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(side)
    section.right_margin = Cm(side)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{SOFTWARE} {VERSION}  {material_name}  第 ")
    set_run_font(r, name="宋体", size=8.2, color=(70, 70, 70))
    add_field(p.add_run(), "PAGE")
    r = p.add_run(" 页")
    set_run_font(r, name="宋体", size=8.2, color=(70, 70, 70))

    footer = section.footer
    footer.paragraphs[0].text = ""
    doc.core_properties.title = f"{SOFTWARE}{VERSION}{material_name}"


def add_blank_applicant_table(doc: Document):
    add_heading(doc, "二、著作权人信息（申请人自行填写）", 1)
    add_notice(doc, "以下内容全部留空。请按实名认证主体和身份证明文件逐字填写；姓名/名称、证件类型和证件号码必须与登记系统及证明文件一致。", "FFF2CC")
    add_table(doc, ["官方字段", "填写栏"], [
        ["姓名或名称", "____________________________________________"],
        ["类别", "□ 自然人  □ 法人  □ 非法人组织  □ 其他"],
        ["证件类型", "____________________________________________"],
        ["证件号码", "____________________________________________"],
        ["国籍", "____________________________________________"],
        ["省份/城市", "____________________________________________"],
        ["联系地址", "____________________________________________"],
        ["邮政编码", "____________________________________________"],
        ["联系人", "____________________________________________"],
        ["手机号码", "____________________________________________"],
        ["电子邮箱", "____________________________________________"],
    ], widths=[1.65, 5.0])


def build_prefill_doc() -> Path:
    doc = Document()
    configure_page(doc, "R11计算机软件著作权登记申请表预填底稿")

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("R11 计算机软件著作权登记申请")
    set_run_font(r, name="微软雅黑", size=20, bold=True, color=(20, 51, 83))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("在线填报预填底稿")
    set_run_font(r, name="微软雅黑", size=15, bold=True, color=(14, 104, 160))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{SOFTWARE}  {VERSION}")
    set_run_font(r, name="微软雅黑", size=12, color=(70, 70, 70))
    doc.add_paragraph()
    add_notice(doc, "重要：这不是中国版权保护中心生成的正式申请表，不能直接上传代替申请确认签章页。请将本稿内容复制到官网登记系统；提交后由系统生成申请确认签章页，必须保持其内容、格式和打印比例不变，打印签章后扫描为 PDF 上传。", "FCE4D6")

    add_heading(doc, "一、软件基本信息", 1)
    add_table(doc, ["官方字段", "预填内容", "提交前核对"], [
        ["软件全称", SOFTWARE, "申请表、说明书和源码页眉必须完全一致。"],
        ["软件简称", "DraARL", "与全称不完全相同，符合官网填写说明。"],
        ["版本号", VERSION, "本材料统一使用 V2.0。提交前应固定 V2.0 源码快照；不要混用 alpha 预发布标识。"],
        ["软件作品说明", "建议选择：原创软件", "初始提交信息含“NRL Link 重构”。须确认前代软件和本次代码均为本人/申报主体自有，且未修改他人软件。"],
        ["修改说明/原登记号", "不适用（若确认原创）", "若属于修改软件，必须改填并按官网要求提交原登记证书或许可证明。"],
        ["开发完成日期", "____年__月__日", "填写 V2.0 全部固定在硬盘等有形载体上的实际日期。"],
        ["发表状态", "建议选择：已发表", "仓库、在线站点和文档已公开；应按实际首次公之于众的事实填写。"],
        ["首次发表日期", "____年__月__日", "Git 首次提交日 2026-03-10 不能自动等同首次公开日。"],
        ["首次发表地点", "中国／互联网（请按系统选项填写）", "首次发表网址可备注：https://github.com/Daofengql/DraARL-Server"],
        ["开发方式", "建议选择：单独开发", "Git 当前仅见单一开发者署名；如存在合作、委托、职务或任务开发，应据实修改并准备证明。"],
    ], widths=[1.35, 2.55, 2.75])

    add_blank_applicant_table(doc)

    add_heading(doc, "三、权利说明", 1)
    add_table(doc, ["官方字段", "预填内容", "核对"], [
        ["权利取得方式", "建议选择：原始取得", "前提是申请软件为申报主体独立开发。"],
        ["继受取得方式", "不适用", "如为受让、继承或承受，应改填并提交对应证明。"],
        ["原登记号", "不适用（如无既有登记）", "已登记软件发生继受时按官网要求填写并提交/寄回相关证书。"],
        ["变更或补充证明编号", "不适用（如无）", "有变更或补充登记时填写。"],
        ["权利范围", "建议选择：全部权利", "如只享有部分权利，须选择部分并注明具体权项。"],
    ], widths=[1.35, 2.55, 2.75])

    add_heading(doc, "四、软件鉴别材料", 1)
    add_table(doc, ["官方字段", "预填内容"], [
        ["交存方式", "一般交存"],
        ["源程序", "《DraARL源程序鉴别材料》：连续前30页 + 连续后30页，共60页，每页50行"],
        ["文档", "《DraARL软件设计说明书鉴别材料》：共60页，每页30行"],
        ["名称和版本", f"两份鉴别材料页眉均为“{SOFTWARE} {VERSION}”，与申请表一致"],
        ["排版", "A4纵向，文字从左向右，页码位于页眉右侧"],
    ], widths=[1.65, 5.0])

    add_heading(doc, "五、软件功能和技术特点", 1)
    add_table(doc, ["官方字段", "可复制的预填内容"], [
        ["开发的硬件环境", "x86-64计算机，4核及以上CPU，8GB及以上内存，SSD存储和网络环境；设备联调可使用ESP32数字电台设备及音频输入输出设备。"],
        ["运行的硬件环境", "x86-64或ARM64服务器，建议2核及以上CPU、4GB及以上内存、SSD和公网/局域网；客户端为具备网络、浏览器和可选麦克风/扬声器的终端。"],
        ["开发的软件环境", "Windows 11、Ubuntu 22.04；Go 1.25、Node.js 20、TypeScript 5.9、React 19、Vite 7、Git、Visual Studio Code；MySQL/MariaDB、Redis、MinIO用于联调。"],
        ["运行的软件环境", "Linux、Windows或macOS；MySQL 5.7+或MariaDB 10.3+；Redis 6.0+可选；MinIO或S3兼容对象存储可选；Chrome、Edge等现代浏览器。"],
        ["编程语言", "Go、TypeScript、JavaScript/TSX、HTML、CSS、SQL"],
        ["源程序量（行数）", "约91,516行（343个生产代码文件，不含测试、第三方依赖、文档和构建产物）"],
        ["开发目的", "为业余无线电及自研数字电台设备提供统一的实时语音/文本通信、设备接入、群组协作、通信记录和平台运维能力。"],
        ["面向领域/行业", "业余无线电、数字通信、物联网设备接入、实时音频通信和通信平台运维管理。"],
        ["软件主要功能", "提供账号注册登录与审核、UDP实体设备和WebSocket客户端接入、动态码绑定、设备参数同步、公开/私有群组、虚拟互联组、PTT语音和文本、多频道接收、通信记录与通联日志、APRS位置、固件/资源发布、管理员后台以及中心/边缘节点互联。"],
        ["软件技术特点", "采用自研DraARLv1协议、按通信域的半双工话权仲裁、按Session去重的接收计划、中心/边缘TLS控制面和HMAC数据面、路由快照与增量投影、短期会话恢复、完整帧过载淘汰、多级缓存及对象存储，实现低时延可扩展的跨设备、跨群组和跨节点通信。"],
    ], widths=[1.65, 5.0])

    add_heading(doc, "六、办理方式、申请人和签章", 1)
    add_table(doc, ["官方字段", "填写栏/预填"], [
        ["申请办理方式", "建议：著作权人申请办理  □ 已确认"],
        ["申请人", "________________________________（应为著作权人之一）"],
        ["代理人", "无／________________________________（委托代理时填写）"],
        ["代理范围、权限和期限", "________________________________（仅委托代理时填写）"],
        ["申请人签章", "由官网生成申请确认签章页后办理；自然人本人签名/人名章，机构加盖鲜章"],
        ["证书领取", "登记完成后在中心官网用户中心下载软件版权登记电子证书"],
    ], widths=[1.75, 4.9])

    add_heading(doc, "七、官网所需文件清单", 1)
    add_bullets(doc, [
        "在线填报申请表后，由官网打印申请确认签章页；不得擅自更改内容、格式和打印比例，签章后上传 PDF 扫描件。",
        "软件鉴别材料：本套件生成的源程序 PDF 和软件设计说明书 PDF。",
        "身份证明文件：根据自然人、企业法人、事业单位、机关、社团、民办非企业或非法人组织类型，按系统提示上传。",
        "其他权属证明：仅在合作、委托、下达任务、修改他人软件、受让、继承或承受等情形下提交。",
        "申请人应留存全部提交文件副本，以便补正时保持内容一致；补正通知通过官网用户中心消息中心送达。",
    ])

    add_heading(doc, "八、官方依据", 1)
    add_para(doc, f"中国版权保护中心《计算机软件著作权登记指南》：{OFFICIAL_GUIDE}")
    add_para(doc, f"《软件著作权登记申请所需文件》：{OFFICIAL_FILES}")
    add_para(doc, f"《计算机软件著作权登记申请表填写说明》：{OFFICIAL_FORM_HELP}")
    add_para(doc, f"在线登记入口：{OFFICIAL_SYSTEM}")
    add_para(doc, f"以上网页核对日期：{ACCESS_DATE}。正式提交时以官网和登记系统当期显示为准。")

    DOCX_OUT.mkdir(parents=True, exist_ok=True)
    path = DOCX_OUT / "01-R11计算机软件著作权登记申请表-预填底稿.docx"
    doc.save(path)
    return path


def clean_markdown_line(line: str) -> str:
    line = line.strip()
    if not line or line.startswith("<!--") or line.startswith("![") or re.search(r"alpha\d*", line, re.IGNORECASE):
        return ""
    if re.fullmatch(r"[|:\-\s]+", line):
        return ""
    line = re.sub(r"^#{1,6}\s*", "", line)
    line = re.sub(r"^>\s*", "", line)
    line = re.sub(r"^[-*+]\s+", "", line)
    line = re.sub(r"^\d+[.)]\s+", "", line)
    line = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", line)
    line = line.replace("`", "")
    if line.startswith("|") and line.endswith("|"):
        cells = [c.strip() for c in line.strip("|").split("|") if c.strip()]
        line = "；".join(cells)
    line = re.sub(r"\s+", " ", line).strip()
    return line


def wrap_document_text(text: str, width=47) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    return textwrap.wrap(text, width=width, break_long_words=True, break_on_hyphens=False, replace_whitespace=True, drop_whitespace=True)


def markdown_to_lines(path: Path) -> list[str]:
    result: list[str] = []
    in_fence = False
    for raw in path.read_text(encoding="utf-8", errors="replace").splitlines():
        if raw.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        cleaned = clean_markdown_line(raw)
        result.extend(wrap_document_text(cleaned))
    return result


def build_manual_line_pool() -> list[str]:
    preface = [
        "第一章 软件概述",
        f"软件名称：{SOFTWARE}",
        f"版本号：{VERSION}",
        "软件简称：DraARL",
        "文档名称：软件设计说明书鉴别材料",
        "DraARL是一套面向业余无线电和自研数字电台设备的实时通信与管理平台。",
        "系统由Go服务端、React浏览器前端、数据存储组件和设备通信协议组成。",
        "服务端同时提供HTTP API、WebSocket在线收发和UDP设备接入服务。",
        "浏览器前端提供公共页面、用户控制台、在线收发页面和管理员后台。",
        "系统支持单机部署，也支持中心节点和多个无数据库边缘节点的分布式部署。",
        "本说明书描述软件用途、功能、架构、协议、数据、安全、部署和运维设计。",
        "本说明书内容依据登记版本源代码、配置模板和仓库技术文档整理。",
        "申请材料不包含第三方依赖源代码、构建产物、密钥、密码和真实用户数据。",
        "第二章 建设目标",
        "为数字电台设备提供统一的身份认证、接入、状态和配置同步能力。",
        "为实体设备、移动客户端、桌面客户端和浏览器提供群组实时通信能力。",
        "为通信管理人员提供用户、设备、群组、节点、资源和记录管理工具。",
        "通过通信域、虚拟互联组和节点投影实现跨群组、跨站点实时转发。",
        "通过权限、限流、加密、审计和会话管理降低非法接入与越权风险。",
        "通过记录、指标、日志、迁移和健康检查支撑持续运行与故障排查。",
        "第三章 总体架构",
        "客户端层包括ESP32硬件设备、浏览器、移动/桌面客户端和软件桥接器。",
        "接入层包括UDP服务器、HTTP服务器、WebSocket服务和TLS节点控制面。",
        "业务层包括认证、设备、群组、路由、话权、记录、资源和审核模块。",
        "数据层包括MySQL或MariaDB、Redis、对象存储及多级内存缓存。",
        "外部服务可包括APRS-IS、Keycloak、SMTP以及可选的OpenAI兼容接口。",
        "中心节点保存权威数据、路由和权限，边缘节点负责设备入口与本地转发。",
        "公开设备协议为DraARLv1，中心与边缘使用受保护的Type 0内部协议。",
        "语音使用Opus 16K数据帧，文本和语音进入统一的通信域路由。",
        "前端资源可独立托管，也可构建后嵌入Go二进制形成单文件服务。",
    ]
    lines: list[str] = []
    for item in preface:
        lines.extend(wrap_document_text(item))

    docs = [
        ROOT / "README.md",
        ROOT / "docs" / "架构设计.md",
        ROOT / "docs" / "usage" / "README.md",
        *sorted((ROOT / "docs" / "usage").glob("*.md")),
        ROOT / "docs" / "Protocol.md",
        ROOT / "docs" / "节点互联协议.md",
        ROOT / "docs" / "数据字典.md",
        ROOT / "docs" / "幽灵多收与消息查询性能验证.md",
        ROOT / "docs" / "UDP单机转发压力测试.md",
        *sorted((ROOT / "docs" / "api").glob("*.md")),
    ]
    seen = set()
    for path in docs:
        if not path.exists() or path.resolve() in seen:
            continue
        seen.add(path.resolve())
        lines.extend(markdown_to_lines(path))

    closing_items = [
        "第二十章 设计总结",
        "DraARL围绕数字无线电设备接入和实时群组通信形成完整的软件功能闭环。",
        "DraARLv1统一定义设备认证、心跳、配置、语音和文本消息的承载方式。",
        "半双工话权仲裁确保同一通信域内同时只有一个有效说话人。",
        "虚拟互联组将多个群组组成通信域并支持多级互联与循环防护。",
        "现代幽灵Session支持同账号多端在线、一个发送频道和多个接收频道。",
        "接收计划按通信域缓存并按Session去重，降低每帧扫描和重复投递成本。",
        "中心与边缘通过TLS控制面同步权威路由并续租设备短期会话。",
        "Type 0数据面通过HMAC、会话世代、时间和重放窗口保护节点中继。",
        "中心重新校验来源所有权和权限，边缘只执行已投影的本地转发计划。",
        "通信记录通过发送时投递快照保存历史可见性，不依赖当前拓扑回算。",
        "音频、固件和客户端资源可以保存在本地或S3兼容对象存储中。",
        "数据库迁移、复合索引、游标分页和查询限流保护大规模历史数据查询。",
        "账号审核、群组权限、设备所有权和管理员角色组成分层授权模型。",
        "JWT、refresh token、设备密码、Origin检查和接口限流构成接入安全边界。",
        "节点凭据轮换、短期会话和失败关闭策略降低长期凭据泄漏影响。",
        "操作日志、通信指标、缓存指标和健康检查支持日常运维和问题定位。",
        "Docker Compose和单二进制部署方式覆盖开发、小规模和集中式部署。",
        "中心/边缘模式用于异地接入、路由域拆分和横向扩展，不改变设备协议。",
        "前端通过React路由区分公共页面、普通用户页面和管理员页面。",
        "设备动态码绑定把设备端请求和已登录用户确认连接为一次性流程。",
        "配置同步支持频率、亚音、静噪、功率和射频保护等设备参数。",
        "在线收发支持PTT、文本、来源频道、发言人显示和连接状态管理。",
        "资源中心、固件OTA和客户端资源分发提供统一的发布管理能力。",
        "APRS位置服务、中继台查询和通联日志扩展业余无线电业务场景。",
        "第三方框架、数据库和编解码库的权利归各自权利人所有。",
        "本登记材料仅主张DraARL自主源代码和软件文档的著作权。",
        "软件名称、版本号和鉴别材料页眉在本套材料中保持一致。",
        "源程序文档和本说明书均采用一般交存方式编排为连续六十页。",
        "正式提交前应固定Git版本、保存源码包并核对所有申请表字段。",
        "正式提交前应确认开发完成日期、首次发表日期和实际著作权归属。",
        "正式提交材料不得包含部署密钥、账号密码、用户数据或未脱敏日志。",
        "系统健康检查接口用于验证HTTP服务是否具备基本响应能力。",
        "UDP入口按照设备身份分片处理，避免同一设备的报文发生乱序。",
        "发送工作线程按照目标稳定分片，尽量保持同一目标的帧顺序。",
        "实时队列超过容量时淘汰最旧完整帧，避免随机截断接收目标。",
        "群组接收者快照在成员变化时失效，并设置定时兜底刷新机制。",
        "设备运行状态与数据库配置分离，减少实时转发路径的数据库访问。",
        "普通实体设备继续保持同一用户和SSID的单端在线约束。",
        "幽灵设备使用客户端实例和临时Session隔离多个在线终端。",
        "发送频道始终包含在接收频道集合中，保证发送端可见本频道状态。",
        "来源频道写入协议保留字段，旧能力节点转发前清零该字段。",
        "路由更新以原子方式替换发送域和接收域，避免部分更新状态。",
        "边缘节点按投影版本确认路由，发现版本缺口时请求重新同步。",
        "路由快照完成长度和摘要校验后一次性提交，避免半成品生效。",
        "中心节点从权威路由重新确定通信域，不信任边缘声明的目标域。",
        "节点数据地址通过TLS挑战和UDP证明绑定，普通数据包不能改绑。",
        "节点控制消息限制帧长和资源用量，防止异常节点占用全部容量。",
        "节点功能位协商用于控制协议扩展，缺少必需能力时拒绝相关会话。",
        "幽灵恢复票据仅在TLS控制面传递，不进入客户端包和公开接口。",
        "恢复后重新查询用户状态和频道权限，不沿用已经撤销的访问权。",
        "数据库记录使用发送者快照，减少用户资料变化对历史记录的影响。",
        "频道消息列表采用时间与主键组合游标，保证分页顺序稳定。",
        "消息查询设置账号、IP和并发限制，避免热点频道拖垮数据库。",
        "对象存储上传采用暂存和完成确认流程，避免未完成对象直接发布。",
        "客户端资源按平台、架构、频道和版本条件生成可下载清单。",
        "固件发布按设备型号和版本规则提供最新版本查询与下载。",
        "APRS连接具备配置、状态和日志管理，用于位置服务和站点信息。",
        "SMTP配置用于验证码和通知邮件，敏感凭据不在公开接口返回。",
        "Keycloak作为可选单点登录服务，不影响本地账号登录能力。",
        "管理员操作通过角色中间件保护，并记录必要的操作审计信息。",
        "用户侧设备操作同时检查登录状态、审核状态和设备所有权。",
        "群主只能管理所属群组范围内的成员、设备和临时通信状态。",
        "公共查询接口使用独立限流规则，降低匿名访问造成的资源竞争。",
        "缓存不可用时关键令牌能力可以降级，但生产环境仍建议启用Redis。",
        "对象存储未启用时可使用本地文件驱动，便于小规模部署和测试。",
        "配置文件提供安全默认值校验，并支持自动生成必要的随机密钥。",
        "数据库仅在显式迁移模式下变更既有结构，降低普通启动风险。",
        "构建流程支持多平台二进制和嵌入式前端资源的发布方式。",
        "文档站持续记录协议、API、数据字典、升级和运维排障信息。",
        "登记版本的最终功能边界以固定源码、配置模板和本说明书为准。",
        "本说明书到此结束。",
    ]
    closing: list[str] = []
    for item in closing_items:
        closing.extend(wrap_document_text(item))
    while len(closing) < 60:
        closing.append(f"附加设计核对项{len(closing) + 1}：登记材料与固定源码保持一致。")
    body = lines[:1740] + closing[-60:]
    if len(body) < 1800:
        raise RuntimeError(f"Insufficient manual lines: {len(body)}")
    return body[:1800]


def add_fixed_line_page(doc: Document, lines: Sequence[str], page_no: int, per_page: int, font_size: float, spacing: float, code=False):
    if page_no > 1:
        doc.add_page_break()
    global_start = (page_no - 1) * per_page
    for offset, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(spacing)
        p.paragraph_format.keep_together = True
        if code:
            text = f"{global_start + offset + 1:05d}  {line}"
            r = p.add_run(text)
            set_run_font(r, name="Consolas", size=font_size, color=(20, 20, 20))
        else:
            text = f"{global_start + offset + 1:04d}  {line}"
            r = p.add_run(text)
            set_run_font(r, name="宋体", size=font_size, color=(25, 25, 25))


def build_manual_identification_doc() -> Path:
    doc = Document()
    configure_identification_page(doc, "软件设计说明书", top=1.4, bottom=1.2, side=1.45)
    lines = build_manual_line_pool()
    for page in range(1, 61):
        chunk = lines[(page - 1) * 30 : page * 30]
        add_fixed_line_page(doc, chunk, page, 30, 8.8, 16.0, code=False)
    path = DOCX_OUT / "02-DraARL软件设计说明书-鉴别材料-60页.docx"
    doc.save(path)
    return path


def production_source_stats() -> tuple[int, int]:
    files: list[Path] = []
    for base in [ROOT / "cmd", ROOT / "internal", ROOT / "pkg", ROOT / "www" / "src"]:
        files.extend(p for p in base.rglob("*") if p.is_file() and p.suffix.lower() in {".go", ".ts", ".tsx", ".css"} and not p.name.endswith("_test.go"))
    total = 0
    for path in files:
        total += len(path.read_text(encoding="utf-8", errors="replace").splitlines())
    return len(files), total


def build_source_identification_doc() -> Path:
    doc = Document()
    configure_identification_page(doc, "源程序", top=1.25, bottom=1.05, side=1.0)
    all_lines = collect_source_lines()
    if len(all_lines) < 3000:
        raise RuntimeError(f"Insufficient source lines: {len(all_lines)}")
    selected = all_lines[:1500] + all_lines[-1500:]
    for page in range(1, 61):
        chunk = selected[(page - 1) * 50 : page * 50]
        add_fixed_line_page(doc, chunk, page, 50, 5.9, 9.2, code=True)
    path = DOCX_OUT / "03-DraARL源程序-鉴别材料-前后各30页.docx"
    doc.save(path)
    return path


def write_official_readme(files: Sequence[Path]):
    file_count, source_lines = production_source_stats()
    text = f"""# DraARL 软件著作权登记材料 - 官方要求版

核对日期：{ACCESS_DATE}

## 这套文件如何使用

1. `01-R11计算机软件著作权登记申请表-预填底稿.docx` 仅用于把项目字段复制到中国版权保护中心登记系统，不能代替系统生成的正式申请表或申请确认签章页。
2. 在线提交后，系统生成申请确认签章页。保持内容、格式和打印比例不变，打印签名/盖章，再扫描为 PDF 上传。
3. `02-DraARL软件设计说明书-鉴别材料-60页.docx/pdf` 为文档鉴别材料，共 60 页，每页 30 行。
4. `03-DraARL源程序-鉴别材料-前后各30页.docx/pdf` 为源程序鉴别材料，共 60 页，每页 50 行，按生产源代码顺序取连续前 1,500 行和连续后 1,500 行。

可编辑 Word 位于 `output/软著材料/官方要求版/`，最终鉴别材料 PDF 位于 `output/pdf/软著材料/`。

## 已统一的申请口径

- 软件全称：{SOFTWARE}
- 软件简称：DraARL
- 版本号：{VERSION}
- 交存方式：一般交存
- 文档类型：软件设计说明书
- 生产源代码统计：{file_count} 个文件，约 {source_lines:,} 行，不含测试、第三方依赖、文档和构建产物

## 申请人自行填写

- 著作权人姓名/名称、类别、证件类型、证件号码、国籍、省份/城市、地址、联系人、电话和邮箱。
- 开发完成日期。
- 首次发表日期；公开仓库和在线站点表明应重点核对“已发表”的实际首次公开日期。
- 确认是否确为原创软件、单独开发、原始取得和全部权利。仓库最早提交写有“NRL Link 重构”，应确认前代代码亦归申报主体所有且没有修改他人软件。

## 官方公开要求摘要

- 申请全程在线办理。
- 所需文件包括申请表、软件鉴别材料和相关证明文件，申请文件纵向排版、文字从左向右。
- 一般交存由源程序和任何一种文档的前、后各连续 30 页组成；整个程序或文档不足 60 页时提交全部。
- 除特定情况外，程序每页不少于 50 行，文档每页不少于 30 行。
- 鉴别材料页眉的软件名称和版本号必须与申请表一致。
- 申请确认签章页必须由系统生成，不得改变内容、格式和打印比例；签章后上传 PDF 扫描件。

## 官方链接

- 登记指南：<{OFFICIAL_GUIDE}>
- 所需文件：<{OFFICIAL_FILES}>
- 填表说明：<{OFFICIAL_FORM_HELP}>
- 在线登记：<{OFFICIAL_SYSTEM}>

正式提交时，以官网和登记系统当期页面为准。
"""
    (DOCX_OUT / "00-使用说明-按中国版权保护中心官网要求.md").write_text(text, encoding="utf-8")

    hashes = []
    for path in files:
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        hashes.append(f"{digest}  {path.name}")
    (DOCX_OUT / "SHA256SUMS.txt").write_text("\n".join(hashes) + "\n", encoding="utf-8")


def main():
    DOCX_OUT.mkdir(parents=True, exist_ok=True)
    PDF_OUT.mkdir(parents=True, exist_ok=True)
    files = [build_prefill_doc(), build_manual_identification_doc(), build_source_identification_doc()]
    write_official_readme(files)
    for path in files:
        print(path)


if __name__ == "__main__":
    main()
